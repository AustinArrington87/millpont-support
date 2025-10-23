# Authored by Austin Arrington May 21, 2025 for MillPont, Inc. 

import pandas as pd
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict
from datetime import datetime

def hectares_to_acres(hectares):
    return hectares * 2.47105

def format_number(num):
    """Format numbers with commas and 2 decimal places"""
    return f"{num:,.2f}"

def set_cell_font(cell, font_name='Open Sans', size=None):
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = font_name
    if size:
        run.font.size = Pt(size)

def add_logo(doc):
    if os.path.exists('logo.png'):
        doc.add_picture('logo.png', width=Inches(2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add space after logo

def set_table_font(table, font_name='Open Sans', size=10):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(size)
                if not paragraph.runs:  # If no runs exist, create one
                    run = paragraph.add_run(paragraph.text)
                    run.font.name = font_name
                    run.font.size = Pt(size)
                    paragraph.text = ''  # Clear original text

def analyze_data(csv_path):
    print(f"Reading CSV file: {csv_path}")
    
    # Read the CSV file
    df = pd.read_csv(csv_path)
    print("CSV loaded successfully")
    
    doc = Document()
    
    # Add logo
    add_logo(doc)
    
    # Add METI™ SSID Registration Report header
    title = doc.add_heading('METI™ SSID Registration Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(204, 0, 0)  # Red color
        run.font.name = 'Open Sans'
    
    # Add report metadata
    metadata = [
        ('Issued By:', 'METI™ Administrators'),
        ('Issued To:', 'Cargill'),
        ('Date:', datetime.now().strftime('%m-%d-%Y'))
    ]
    
    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(f"{label} {value}").font.name = 'Open Sans'
    
    doc.add_paragraph()
    
    # Add report introduction
    intro = doc.add_paragraph()
    intro_run = intro.add_run("This report has been prepared at the request of Cargill, a METI™ Member. It provides verification of the registration, uniqueness, and conflict status of digital deeds called Secure Source IDs (SSIDs) on the MillPont Environmental Trust Infrastructure (METI™) platform.")
    intro_run.font.name = 'Open Sans'
    
    platform_desc = doc.add_paragraph()
    platform_run = platform_desc.add_run("The METI™ platform operates a geospatial clearinghouse to register and verify SSIDs, ensuring the integrity and exclusivity of environmental asset boundaries. Leveraging SSIDs registered on the METI™ Source Ledger, METI™ ensures data integrity, transparency, and conflict detection and resolution for all registered Sources of environmental assets in agricultural supply chains.")
    platform_run.font.name = 'Open Sans'
    
    doc.add_paragraph()
    
    # Add Verification Scope section
    doc.add_heading('Verification Scope', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    scope = doc.add_paragraph()
    scope_run = scope.add_run("The scope of this report includes the verification of Secure Source IDs registered by Cargill on the METI™ platform. It addresses the following aspects:")
    scope_run.font.name = 'Open Sans'
    
    scope_items = [
        "Registration: Confirmation that the Source IDs have been successfully registered within the METI™ Source Ledger.",
        "Uniqueness: Assurance that the registered Source IDs are unique and have not been duplicated within the platform.",
        "Conflict Status: Current status of conflict checks performed on the registered Secure Source IDs, including any status of resolutions or pending disputes."
    ]
    
    for item in scope_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item).font.name = 'Open Sans'
    
    note = doc.add_paragraph()
    note_run = note.add_run("It is important to note that this report reflects a moment in time within a dynamic and continually updating system. As new data is processed and reconciliations occur, conflict statuses and other verifications may evolve.")
    note_run.font.name = 'Open Sans'
    
    commitment = doc.add_paragraph()
    commitment_run = commitment.add_run("This verification aligns with METI™'s commitment to providing transparent and robust data assurance for its members and their stakeholders.")
    commitment_run.font.name = 'Open Sans'
    
    doc.add_paragraph()

    # Add Verification Details section
    doc.add_heading('Verification Details', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    details = [
        ('Member Name:', 'Cargill'),
        ('Number of Secure Source IDs Registered:', f"{len(df):,}"),
        ('Unique Secure Source IDs (no potential conflict detected):', f"{len(df[df['conflict'] == False]):,}"),
        ('Potential Conflicts Detected (pending resolution):', f"{len(df[df['conflict'] == True]):,} (~{len(df[df['conflict'] == True])/len(df)*100:.1f}%)")
    ]
    
    for label, value in details:
        p = doc.add_paragraph()
        p.add_run(f"{label} {value}").font.name = 'Open Sans'
    
    doc.add_paragraph()
    
    # Add Analysis section
    doc.add_heading('Analysis', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    analysis = doc.add_paragraph()
    analysis_run = analysis.add_run(f"As part of routine verification, the system flagged {len(df[df['conflict'] == True]):,} of {len(df):,} SSIDs for \"potential conflicts.\" These potential conflicts are indications, not conclusions, and require further review by the Member, affiliated data stakeholders and METI™ Administration. They can arise from several scenarios, including:")
    analysis_run.font.name = 'Open Sans'
    
    conflict_scenarios = [
        "Duplicate Entries: Instances where duplicate data for registering different SSIDs was submitted due to human or system errors in data entry.",
        "Boundary Overlaps: Overlaps may occur when spatial boundaries registered for different SSIDs intersect partially. These overlaps could result from measurement inaccuracies, data formatting variations, or the inherent complexities in delineating irregular landscapes.",
        "Data Representation Inaccuracies: Geospatial data quality issues, such as outdated maps, varying coordinate systems, or errors during digitization, can contribute to discrepancies flagged by the system."
    ]
    
    for scenario in conflict_scenarios:
        p = doc.add_paragraph(style='List Number')
        p.add_run(scenario).font.name = 'Open Sans'
    
    doc.add_paragraph()
    
    # Add METI™ Platform Assurance section
    doc.add_heading('METI™ Platform Assurance & Administrator Declaration', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    assurance = doc.add_paragraph()
    assurance_run = assurance.add_run("The METI™ platform employs rigorous standards for ensuring the validity and integrity of registered data, including:")
    assurance_run.font.name = 'Open Sans'
    
    standards = [
        "Unique ID generation and cross-checking to prevent duplication, including ongoing monitoring until claim expiration.",
        "Automated and manual conflict resolution processes.",
        "Secure storage and encryption of all data."
    ]
    
    for standard in standards:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(standard).font.name = 'Open Sans'
    
    declaration = doc.add_paragraph()
    declaration_run = declaration.add_run(f"This report affirms that the Secure Source IDs listed under Cargill meet the platform's standards for registration, uniqueness, and conflict-free status, except as noted in the conflict details above.")
    declaration_run.font.name = 'Open Sans'
    
    # Add Signature section
    doc.add_heading('Signature Page', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    administrators = [
        ('Name:', 'Dominic Sutton-Vermeulen', 'Title:', 'CEO'),
        ('Name:', 'Austin Arrington', 'Title:', 'CTO')
    ]
    
    for admin in administrators:
        p = doc.add_paragraph()
        p.add_run(f"{admin[0]} {admin[1]}\n{admin[2]} {admin[3]}\nSignature:\n___________________________\nDate: {datetime.now().strftime('%m/%d/%Y')}").font.name = 'Open Sans'
        doc.add_paragraph()

    # Add title with red color
    title = doc.add_heading('Field Boundary Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(204, 0, 0)  # Red color
        run.font.name = 'Open Sans'
    
    # Add custodian
    custodian = doc.add_paragraph()
    custodian.alignment = WD_ALIGN_PARAGRAPH.CENTER
    custodian_run = custodian.add_run('Custodian: Cargill')
    custodian_run.font.name = 'Open Sans'
    custodian_run.bold = True
    
    # Add timestamp
    timestamp = doc.add_paragraph()
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_run = timestamp.add_run(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp_run.font.name = 'Open Sans'
    
    doc.add_paragraph()

    # Description
    doc.add_heading('Description', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    desc = doc.add_paragraph()
    desc_run = desc.add_run("This report analyzes field boundaries' conflicts with other boundaries in the METI™ Source Ledger, as well as overlaps with protected areas.")
    desc_run.font.name = 'Open Sans'
    
    doc.add_paragraph()

    # Overall Summary
    doc.add_heading('Overall Summary', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    total_sources = len(df)
    total_hectares = df['hectares'].sum()
    total_acres = hectares_to_acres(total_hectares)
    
    # Get conflict and protected area statistics
    sources_with_conflicts = len(df[df['conflict'] == True])
    sources_with_protected = len(df[df['unep_overlap'] == True])
    
    # Get internal vs external conflict breakdown
    internal_conflicts = len(df[df['is_internal'] == True])
    external_conflicts = sources_with_conflicts - internal_conflicts
    
    # Get internal conflicts with >=2% overlap
    internal_conflicts_df = df[df['is_internal'] == True]
    high_overlap_conflicts = internal_conflicts_df[internal_conflicts_df['percent_overlap'] >= 0.02]
    
    summary = doc.add_paragraph()
    summary_run = summary.add_run('Total Sources: ')
    summary_run.bold = True
    summary_run.font.name = 'Open Sans'
    summary.add_run(f"{total_sources:,}\n").font.name = 'Open Sans'
    
    summary_run = summary.add_run('Total Area: ')
    summary_run.bold = True
    summary_run.font.name = 'Open Sans'
    summary.add_run(f"{format_number(total_hectares)} hectares ({format_number(total_acres)} acres)\n").font.name = 'Open Sans'
    
    summary_run = summary.add_run('Sources With Conflicts: ')
    summary_run.bold = True
    summary_run.font.name = 'Open Sans'
    summary.add_run(f"{sources_with_conflicts:,} ({(sources_with_conflicts/total_sources)*100:.1f}%)\n").font.name = 'Open Sans'
    
    summary_run = summary.add_run('  - Internal Conflicts: ')
    summary_run.bold = True
    summary_run.font.name = 'Open Sans'
    summary.add_run(f"{internal_conflicts:,} ({(internal_conflicts/total_sources)*100:.1f}%)\n").font.name = 'Open Sans'
    
    summary_run = summary.add_run('  - External Conflicts: ')
    summary_run.bold = True
    summary_run.font.name = 'Open Sans'
    summary.add_run(f"{external_conflicts:,} ({(external_conflicts/total_sources)*100:.1f}%)\n").font.name = 'Open Sans'
    
    summary_run = summary.add_run('Internal Conflicts with ≥2% Overlap: ')
    summary_run.bold = True
    summary_run.font.name = 'Open Sans'
    summary.add_run(f"{len(high_overlap_conflicts):,} ({(len(high_overlap_conflicts)/total_sources)*100:.1f}%)\n").font.name = 'Open Sans'
    
    summary_run = summary.add_run('Sources With Protected Area Overlaps: ')
    summary_run.bold = True
    summary_run.font.name = 'Open Sans'
    summary.add_run(f"{sources_with_protected:,} ({(sources_with_protected/total_sources)*100:.1f}%)").font.name = 'Open Sans'

    # Protected Areas Analysis
    doc.add_heading('Protected Areas Analysis', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    # Create table for protected areas overlaps
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    headers = ['Total Sources', 'Protected Area Overlaps', 'Overlap Percentage']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    row_cells = table.add_row().cells
    row_cells[0].text = f"{total_sources:,}"
    row_cells[1].text = f"{sources_with_protected:,}"
    row_cells[2].text = f"{(sources_with_protected/total_sources)*100:.1f}%"
    
    set_table_font(table)
    doc.add_paragraph()

    # High Overlap Conflicts Analysis
    if len(high_overlap_conflicts) > 0:
        doc.add_heading('High Overlap Conflict Analysis', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
        
        # Add client SOP note
        sop_note = doc.add_paragraph()
        sop_run = sop_note.add_run("The following fields show overlap greater than 2% and per the client's SOP, a deduction should be made from the outcome.")
        sop_run.font.name = 'Open Sans'
        sop_run.bold = True
        
        doc.add_paragraph()
        
        # Create table for high overlap conflicts
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = ['Source ID', 'Alt ID', 'Country', 'Overlap %', 'Area (Hectares)']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add high overlap conflict details
        for _, row in high_overlap_conflicts.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['id'])
            row_cells[1].text = str(row['alt_id']) if pd.notna(row['alt_id']) else '-'
            row_cells[2].text = str(row['country']) if pd.notna(row['country']) else '-'
            row_cells[3].text = f"{row['percent_overlap']*100:.2f}%"
            row_cells[4].text = format_number(row['hectares'])
        
        set_table_font(table)
        doc.add_paragraph()

    # External Overlap Conflict Analysis
    external_conflicts_df = df[df['is_internal'] == False]
    if len(external_conflicts_df) > 0:
        doc.add_heading('External Overlap Conflict Analysis', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
        
        # Add explanatory text
        external_note = doc.add_paragraph()
        external_run = external_note.add_run("Spatiotemporal overlaps were found with sources managed by other METI™ custodians. Supporting documentation has been requested by all parties and METI™ Administrators are working towards resolution.")
        external_run.font.name = 'Open Sans'
        
        doc.add_paragraph()
        
        # Create table for external conflicts
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = ['Source ID', 'Alt ID', 'Country', 'Overlap %', 'Area (Hectares)']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add external conflict details
        for _, row in external_conflicts_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['id'])
            row_cells[1].text = str(row['alt_id']) if pd.notna(row['alt_id']) else '-'
            row_cells[2].text = str(row['country']) if pd.notna(row['country']) else '-'
            row_cells[3].text = f"{row['percent_overlap']*100:.2f}%"
            row_cells[4].text = format_number(row['hectares'])
        
        set_table_font(table)
        doc.add_paragraph()

    # Protected Areas Analysis by Country
    doc.add_heading('Protected Areas Analysis by Country', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    for country in sorted(df['country'].unique(), key=lambda x: str(x) if pd.notna(x) else ''):
        if pd.isna(country):
            continue
            
        # Filter out rows where both pa_name and pa_designation are NaN
        country_pas = df[df['country'] == country][['pa_name', 'pa_designation']]
        country_pas = country_pas.dropna(how='all')  # Remove rows where all values are NaN
        country_pas = country_pas.drop_duplicates()
        
        # Only create section if country has protected areas
        if len(country_pas) > 0:
            doc.add_heading(f'Protected Areas in {country}', level=2).runs[0].font.color.rgb = RGBColor(204, 0, 0)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            headers = ['Protected Area Name', 'Designation']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for _, row in country_pas.iterrows():
                if pd.notna(row['pa_name']) or pd.notna(row['pa_designation']):  # Only add row if at least one value is not NaN
                    row_cells = table.add_row().cells
                    for i, value in enumerate([
                        str(row['pa_name']) if pd.notna(row['pa_name']) else '-', 
                        str(row['pa_designation']) if pd.notna(row['pa_designation']) else '-'
                    ]):
                        row_cells[i].text = value
            
            set_table_font(table)
            doc.add_paragraph()

    # Source Details with Protected Area Overlaps
    doc.add_heading('Source Details with Protected Area Overlaps', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    # Filter for sources with protected area overlaps
    protected_sources = df[df['unep_overlap'] == True]
    
    for country in sorted(protected_sources['country'].unique(), key=lambda x: str(x) if pd.notna(x) else ''):
        if pd.isna(country):
            continue
            
        country_sources = protected_sources[protected_sources['country'] == country]
        if len(country_sources) > 0:
            doc.add_heading(f'Country: {country}', level=2).runs[0].font.color.rgb = RGBColor(204, 0, 0)
            
            # Create table for source details
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            headers = ['Source ID', 'Alt ID', 'Protected Area', 'Area (Hectares)']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Add source details
            for _, row in country_sources.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['id'])
                row_cells[1].text = str(row['alt_id']) if pd.notna(row['alt_id']) else '-'
                row_cells[2].text = str(row['pa_name']) if pd.notna(row['pa_name']) else '-'
                row_cells[3].text = format_number(row['hectares'])
            
            set_table_font(table)
            doc.add_paragraph()

    # Country Analysis
    doc.add_heading('Country Analysis', level=1).runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    # Get country-level statistics
    for country in sorted(df['country'].unique(), key=lambda x: str(x) if pd.notna(x) else ''):
        if pd.isna(country):
            continue
        
        # Get country data
        country_data = df[df['country'] == country]
        total_country_sources = len(country_data)
        country_conflicts = len(country_data[country_data['conflict'] == True])
        country_internal_conflicts = len(country_data[country_data['is_internal'] == True])
        country_external_conflicts = country_conflicts - country_internal_conflicts
        country_protected = len(country_data[country_data['unep_overlap'] == True])
        
        if total_country_sources > 0:
            doc.add_heading(f'Country: {country}', level=2).runs[0].font.color.rgb = RGBColor(204, 0, 0)
            
            # Create table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Set headers
            headers = ['Total Sources', 'Total Conflicts', 'Internal Conflicts', 'External Conflicts', 'Sources With Protected Areas', 'Total Area (Hectares)']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            
            # Add row
            row_cells = table.add_row().cells
            row_cells[0].text = f"{total_country_sources:,}"
            row_cells[1].text = f"{country_conflicts:,} ({(country_conflicts/total_country_sources)*100:.1f}%)"
            row_cells[2].text = f"{country_internal_conflicts:,} ({(country_internal_conflicts/total_country_sources)*100:.1f}%)"
            row_cells[3].text = f"{country_external_conflicts:,} ({(country_external_conflicts/total_country_sources)*100:.1f}%)"
            row_cells[4].text = f"{country_protected:,} ({(country_protected/total_country_sources)*100:.1f}%)"
            row_cells[5].text = format_number(country_data['hectares'].sum())
            
            set_table_font(table)
            doc.add_paragraph()

    # Save the document
    output_filename = 'meti_ssid_registration_report.docx'
    doc.save(output_filename)
    print(f"Report generated: {output_filename}")

if __name__ == "__main__":
    csv_file = "sources_20251023_145659.csv"
    
    if not os.path.exists(csv_file):
        print(f"Error: Could not find {csv_file}")
    else:
        analyze_data(csv_file) 
